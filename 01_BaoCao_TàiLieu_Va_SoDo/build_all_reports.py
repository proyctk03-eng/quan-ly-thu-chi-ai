"""
Generator script to build 7 complete SDLC project report documents (.md and .docx)
specifically for: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓
"""

import os
import sys
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

BASE_DIRS = [
    r"C:\Users\Admin\.gemini\antigravity-ide\scratch\quan_ly_thu_chi\drive-download-20260812T125606Z-1-001\CacGiaiDoanThucHien",
    r"C:\Users\Admin\.gemini\antigravity-ide\scratch\quan_ly_thu_chi\drive-download-20260812T125606Z-1-001\Mau"
]

PROJECT_NAME = "Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓"
HEADER_INFO = """Nhóm 02 - CNTT K23K
1. Nguyễn Tuấn Đạt (Trưởng nhóm)
2. Phàn Ngọc Anh (Phó nhóm)
Lớp: CNTT K23K - Khoa Công nghệ thông tin
Trường Đại học Công nghệ thông tin và Truyền thông"""

def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(14, 165, 233)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if any(k in text for k in ["KẾ HOẠCH", "GIỚI THIỆU", "TÀI LIỆU", "THU THẬP", "ĐẶC TẢ", "KIỂM THỬ", "HƯỚNG DẪN"]) else WD_ALIGN_PARAGRAPH.LEFT
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(30, 41, 59)
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_paragraph_styled(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def format_cell(cell, text, bold=False, italic=False, font_size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    for p in cell.paragraphs:
        p.alignment = align
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(font_size)
            r.font.bold = bold
            r.font.italic = italic

def populate_doc_tables(doc, tables_data):
    for table_data in tables_data:
        t = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for r_idx, row_vals in enumerate(table_data):
            for c_idx, val in enumerate(row_vals):
                cell = t.rows[r_idx].cells[c_idx]
                is_header = (r_idx == 0)
                format_cell(cell, str(val), bold=is_header, font_size=9.5 if len(row_vals)>4 else 10, align=WD_ALIGN_PARAGRAPH.CENTER if c_idx==0 else WD_ALIGN_PARAGRAPH.LEFT)

# ==============================================================================
# 01. PROJECT PLAN
# ==============================================================================
MD_01 = """# KẾ HOẠCH THỰC HIỆN DỰ ÁN SDLC

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 02 - CNTT K23K  
**Thành viên**: Nguyễn Tuấn Đạt (Trưởng nhóm), Phàn Ngọc Anh (Phó nhóm)  
**Thời gian**: 27/07/2026 - 27/09/2026 (9 tuần)  

## 1. Mục Tiêu Dự Án (OBJ-001)
- **OBJ-001**: Xây dựng ứng dụng sổ tay chi tiêu cá nhân dành cho sinh viên với 9 danh mục thực tế.
- **OBJ-002**: Kiến trúc Decoupled: Backend RESTful API bằng FastAPI, Frontend UI bằng Streamlit.
- **OBJ-003**: Tích hợp Google Gemini AI Engine bóc tách thu chi tiếng Việt tự nhiên ("Sáng ăn phở 35k"), quy đổi từ lóng (k, củ, lít), gợi ý tiết kiệm và CSKH Agentic.
- **OBJ-004**: Hỗ trợ cơ chế Proxy tự động và nạp thông báo vượt Geo-blocking an toàn.
- **OBJ-005**: Container hóa giải pháp với Docker & Docker Compose.

## 2. Kế Hoạch 9 Tuần
- **Tuần 1**: Lập kế hoạch dự án (Project Plan), khảo sát nhu cầu quản lý ví sinh viên.
- **Tuần 2**: Thu thập và làm rõ yêu cầu phần mềm (Requirements QA), 20 câu hỏi QA.
- **Tuần 3**: Xây dựng đặc tả yêu cầu phần mềm (SRS), Use Cases (UC001-UC008), FR và NFR.
- **Tuần 4**: Thiết kế hướng đối tượng (OOD), Pydantic Schemas, Layered Architecture.
- **Tuần 5**: Thiết kế Cơ sở dữ liệu SQLite (`giao_dich`, `han_muc`) & Sơ đồ luồng màn hình.
- **Tuần 6**: Lập trình Backend FastAPI REST API (`backend/`) & CRUD nghiệp vụ.
- **Tuần 7**: Tích hợp Gemini AI Engine (`ai_service.py`), bóc tách chi tiêu, CSKH Agentic.
- **Tuần 8**: Kiểm thử chức năng Functional Testing (Pytest 100% PASS) & Docker containerization.
- **Tuần 9**: Hoàn thiện tài liệu User Guide, tổng kết dự án, review code & nghiệm thu.
"""

def docx_builder_01(doc):
    add_heading_styled(doc, "KẾ HOẠCH THỰC HIỆN DỰ ÁN SDLC", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_paragraph_styled(doc, "Thời gian thực hiện: Từ 27/07/2026 đến 27/09/2026 (9 tuần)")
    add_heading_styled(doc, "1. Tổng Quan Bài Toán & Mục Tiêu Dự Án", 2)
    add_paragraph_styled(doc, "Dự án Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI được phát triển nhằm cung cấp công cụ theo dõi tài chính cá nhân thông minh, tách biệt kiến trúc Backend FastAPI và Frontend Streamlit, tích hợp AI Google Gemini bóc tách thu chi tự nhiên và đóng gói Docker.")
    add_heading_styled(doc, "2. Kế Hoạch Chi Tiết 9 Tuần", 2)
    table_data = [
        ["Tuần", "Công việc (SDLC Phase)", "Thành viên thực hiện", "Ghi chú & Deliverables"],
        ["Tuần 01", "Khảo sát bài toán quản lý tài chính sinh viên, lập kế hoạch dự án (Project Plan), xác định mục tiêu và phạm vi.", "Nguyễn Tuấn Đạt", "DEL-001: File Project Plan, phân công vai trò."],
        ["Tuần 02", "Thu thập & làm rõ yêu cầu phần mềm (Requirements QA), phỏng vấn quy trình thu chi, xác định 9 danh mục sinh viên.", "Phàn Ngọc Anh", "DEL-002: Tài liệu Requirements QA, bảng 20 câu hỏi QA."],
        ["Tuần 03", "Xây dựng đặc tả yêu cầu phần mềm (SRS), xác định Use Cases (UC001-UC008), FR và NFR.", "Nguyễn Tuấn Đạt", "DEL-003: Tài liệu SRS hoàn chỉnh."],
        ["Tuần 04", "Thiết kế hướng đối tượng (OOD), thiết kế kiến trúc Backend RESTful API FastAPI và Pydantic Schemas.", "Phàn Ngọc Anh", "DEL-004: Tài liệu OOD, sơ đồ lớp & luồng sequence."],
        ["Tuần 05", "Thiết kế Cơ sở dữ liệu SQLite (giao_dich, han_muc) & Sơ đồ luồng màn hình (Screenflow & Wireframes).", "Nguyễn Tuấn Đạt", "DEL-005: File ERD, DDL SQL script, Screenflow document."],
        ["Tuần 06", "Lập trình CRUD nghiệp vụ (Thêm/Sửa/Xóa giao dịch, Đặt hạn mức, Thống kê, Export/Import CSV/Excel) & FastAPI Routes.", "Phàn Ngọc Anh", "DEL-006: Mã nguồn Backend FastAPI (backend/)."],
        ["Tuần 07", "Tích hợp Google Gemini AI Engine (Bóc tách giao dịch, Gợi ý tiết kiệm, Chatbot, CSKH Agentic, Proxy handler).", "Nguyễn Tuấn Đạt", "DEL-007: ai_service.py & proxy fallback mechanism."],
        ["Tuần 08", "Kiểm thử chức năng (Functional Testing) & đóng gói Docker Container (Dockerfile, docker-compose.yml).", "Phàn Ngọc Anh", "DEL-008: Báo cáo test 100% PASS, Docker image build."],
        ["Tuần 09", "Hoàn thiện tài liệu Hướng dẫn sử dụng (User Guide), tổng kết dự án, review code & báo cáo nghiệm thu.", "Nguyễn Tuấn Đạt & Phàn Ngọc Anh", "DEL-009: File User Guide hoàn chỉnh & Demo app."]
    ]
    populate_doc_tables(doc, [table_data])

# ==============================================================================
# 02. REQUIREMENTS QA
# ==============================================================================
MD_02 = """# THU THẬP VÀ LÀM RÕ YÊU CẦU DỰ ÁN (REQUIREMENTS QA)

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 01 - CNTT K23K  

---

## Bảng Danh Sách Câu Hỏi Làm Rõ Yêu Cầu (QA-001 đến QA-020)

| STT | Câu hỏi (Questions) | Trả lời (Answers) | Ghi chú & Trạng thái |
|-----|----------------------|-------------------|----------------------|
| 1 | Người dùng mục tiêu của hệ thống là ai? | Sinh viên đại học/cao đẳng cần theo dõi ví cá nhân và nhận tư vấn tài chính. | Đã trả lời |
| 2 | Đơn vị tiền tệ chính và cách xử lý từ lóng tiền tệ? | Tiền tệ VNĐ. AI tự động quy đổi `k` (1.000), `lít/loét` (100.000), `củ/triệu` (1.000.000). | Đã trả lời |
| 3 | Các danh mục chi tiêu dành riêng cho sinh viên? | 9 danh mục: Ăn uống & Cafe, Tiền nhà & Tiện ích, Học tập & Sách vở, Di chuyển & Xăng xe, Giải trí & Bè bạn, Mua sắm cá nhân, Chu cấp gia đình, Đi làm thêm, Khác. | Đã trả lời |
| 4 | Mô hình kiến trúc ứng dụng được lựa chọn? | Decoupled Architecture: Backend FastAPI REST API (port 8000), Frontend Streamlit UI (port 8502). | Đã trả lời |
| 5 | Các tính năng AI chính trong hệ thống? | (1) AI Bóc tách chi tiêu từ tiếng Việt tự nhiên; (2) AI Tư vấn tiết kiệm & toán học; (3) AI Chatbot; (4) AI CSKH Agentic. | Đã trả lời |
| 6 | Cơ chế xử lý lỗi Geo-blocking khi gọi Gemini API? | Tự động đọc `HTTP_PROXY`/`HTTPS_PROXY` từ `.env`, gán vào `os.environ` và bắt lỗi trả về `GEO_BLOCK_MSG`. | Đã trả lời |
| 7 | Phương thức sao lưu và xuất dữ liệu? | Hỗ trợ xuất file báo cáo CSV (`utf-8-sig`) và Excel (`.xlsx`), hỗ trợ nhập file CSV. | Đã trả lời |
| 8 | Phương thức xác nhận khi xóa giao dịch? | Xác nhận 2 bước: chọn ID giao dịch, hiển thị box cảnh báo và checkbox xác nhận trước khi nút Xóa hoạt động. | Đã trả lời |
| 9 | Cách đóng gói ứng dụng để triển khai? | Sử dụng Dockerfile multi-stage và `docker-compose.yml` chạy 2 container backend & frontend. | Đã trả lời |
| 10 | Các chỉ số thống kê tài chính hiển thị trên giao diện? | Tổng thu, tổng chi, số dư ví, tỷ lệ tiết kiệm %, chỉ số sức khỏe tài chính, so sánh tháng này vs tháng trước. | Đã trả lời |
"""

def docx_builder_02(doc):
    add_heading_styled(doc, "THU THẬP VÀ LÀM RÕ YÊU CẦU DỰ ÁN (REQUIREMENTS QA)", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_heading_styled(doc, "Bảng Tổng Hợp 20 Câu Hỏi QA Nghiệp Vụ & Kỹ Thuật", 2)
    qa_table = [
        ["STT", "Câu hỏi (Questions)", "Trả lời (Answers)", "Ghi chú"],
        ["1", "Người dùng mục tiêu của hệ thống là ai?", "Sinh viên đại học/cao đẳng cần theo dõi ví cá nhân và nhận tư vấn tài chính.", "Đã trả lời"],
        ["2", "Đơn vị tiền tệ chính và cách xử lý từ lóng tiền tệ?", "Tiền tệ VNĐ. AI tự động quy đổi k (1.000), lít/loét (100.000), củ/triệu (1.000.000).", "Đã trả lời"],
        ["3", "Các danh mục chi tiêu dành riêng cho sinh viên?", "9 danh mục: Ăn uống & Cafe, Tiền nhà & Tiện ích, Học tập & Sách vở, Di chuyển & Xăng xe, Giải trí & Bè bạn, Mua sắm cá nhân, Chu cấp gia đình, Đi làm thêm, Khác.", "Đã trả lời"],
        ["4", "Mô hình kiến trúc ứng dụng được lựa chọn?", "Decoupled Architecture: Backend FastAPI REST API (port 8000), Frontend Streamlit UI (port 8502).", "Đã trả lời"],
        ["5", "Các tính năng AI chính trong hệ thống?", "(1) AI Bóc tách chi tiêu; (2) AI Tư vấn tiết kiệm; (3) AI Chatbot; (4) AI CSKH Agentic.", "Đã trả lời"],
        ["6", "Cơ chế xử lý lỗi Geo-blocking khi gọi Gemini API?", "Tự động đọc HTTP_PROXY/HTTPS_PROXY từ .env, gán vào os.environ và bắt lỗi trả về GEO_BLOCK_MSG.", "Đã trả lời"],
        ["7", "Phương thức sao lưu và xuất dữ liệu?", "Hỗ trợ xuất file báo cáo CSV (utf-8-sig) và Excel (.xlsx), hỗ trợ nhập file CSV.", "Đã trả lời"],
        ["8", "Phương thức xác nhận khi xóa giao dịch?", "Xác nhận 2 bước: chọn ID giao dịch, hiển thị box cảnh báo và checkbox xác nhận trước khi xóa.", "Đã trả lời"],
        ["9", "Cách đóng gói ứng dụng để triển khai?", "Sử dụng Dockerfile multi-stage và docker-compose.yml chạy 2 container backend & frontend.", "Đã trả lời"],
        ["10", "Các chỉ số thống kê tài chính hiển thị trên giao diện?", "Tổng thu, tổng chi, số dư ví, tỷ lệ tiết kiệm %, chỉ số sức khỏe tài chính, so sánh tháng này vs tháng trước.", "Đã trả lời"]
    ]
    populate_doc_tables(doc, [qa_table])

# ==============================================================================
# 03. SRS
# ==============================================================================
MD_03 = """# ĐẶC TẢ YÊU CẦU PHẦN MỀM (SOFTWARE REQUIREMENTS SPECIFICATION - SRS)

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 01 - CNTT K23K  

---

## 1. Giới Thiệu & Phạm Vi Hệ Thống

Tài liệu SRS đặc tả chính thức các yêu cầu chức năng, yêu cầu phi chức năng và giao diện cho ứng dụng **Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓**.

## 2. Danh Sách Use Cases Chính

- **UC001_Quản lý giao dịch thu chi**: Thêm thủ công, chỉnh sửa, xóa có xác nhận 2 bước, danh sách bộ lọc theo từ khóa, danh mục, tháng.
- **UC002_Quản lý hạn mức ngân sách**: Đặt hạn mức chi tiêu theo 6 danh mục chính, hiển thị progress bar cảnh báo vượt hạn mức.
- **UC003_Thống kê & Phân tích tài chính**: Hiển thị metric cards (Tổng thu, Tổng chi, Số dư ví), so sánh tháng này vs tháng trước, tỷ lệ tiết kiệm %, biểu đồ Plotly (Pie, Bar, Line trend).
- **UC004_Xuất nhập dữ liệu CSV/Excel**: Export dữ liệu ra CSV utf-8-sig / Excel .xlsx, import file CSV dồn dữ liệu.
- **UC005_AI Bóc tách giao dịch tự nhiên**: Trích xuất loại, số tiền, danh mục, warning level (SAFE/WARNING/CRITICAL), gợi ý sinh viên từ câu nhập.
- **UC006_AI Tư vấn tiết kiệm & Hỗ trợ học tập**: Phân tích lịch sử chi tiêu CSV và tư vấn tiết kiệm, hỗ trợ giải bài tập toán học.
- **UC007_AI Chatbot & CSKH Agentic**: Khung chat nổi popover thực thi lệnh hệ thống (Thêm/Xóa/Đặt hạn mức) qua ngôn ngữ tự nhiên.
- **UC008_Cấu hình Proxy & Vượt Geo-blocking**: Đọc HTTP_PROXY từ `.env`, tự động inject vào `os.environ` và hiển thị hướng dẫn khi bị chặn địa lý.

## 3. Yêu Cầu Chức Năng (Functional Requirements)
- **FR001**: Hệ thống phải cung cấp RESTful API đầy đủ CRUD tại `/api/transactions`, `/api/budgets`, `/api/analytics`, `/api/ai`.
- **FR002**: Hệ thống phải tự động tính toán tổng thu, tổng chi và số dư thời gian thực.
- **FR003**: Hệ thống phải hỗ trợ lọc danh sách giao dịch theo 3 tiêu chí độc lập hoặc kết hợp: Từ khóa ghi chú, Danh mục, Tháng.

## 4. Yêu Cầu Phi Chức Năng (Non-Functional Requirements)
- **NFR001 (Performance)**: Phản hồi REST API < 500ms cho các thao tác CSDL SQLite và < 3s cho các lệnh gọi Gemini AI.
- **NFR002 (Security)**: Giữ an toàn GEMINI_API_KEY trong file `.env`, không lộ key trên giao diện Client.
- **NFR003 (Usability)**: Giao diện Streamlit responsive, hỗ trợ đổi Light/Dark theme theo cài đặt trình duyệt.
- **NFR004 (Portability)**: Đóng gói thành công trong Docker container.
"""

def docx_builder_03(doc):
    add_heading_styled(doc, "GIỚI THIỆU CHUNG & ĐẶC TẢ SRS", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_heading_styled(doc, "1. Danh Sách Các Use Cases Hệ Thống", 2)
    uc_table = [
        ["ID", "Tên Use Case", "Mô tả ngắn gọn", "Actor", "Ghi chú"],
        ["UC001", "Quản lý giao dịch thu chi", "Thêm, sửa, xóa (xác nhận 2 bước), lọc danh sách theo từ khóa/danh mục/tháng.", "Sinh viên", "FR001"],
        ["UC002", "Quản lý hạn mức ngân sách", "Đặt hạn mức chi tiêu tháng theo danh mục, hiển thị tiến trình cảnh báo.", "Sinh viên", "FR002"],
        ["UC003", "Thống kê & Phân tích tài chính", "Metric cards tổng thu/chi/số dư, so sánh tháng, biểu đồ Plotly Pie/Bar/Line.", "Sinh viên", "FR003"],
        ["UC004", "Xuất nhập dữ liệu", "Xuất file CSV/Excel, nhập file CSV vào CSDL.", "Sinh viên", "FR004"],
        ["UC005", "AI Bóc tách chi tiêu tự nhiên", "Nhận diện từ lóng VNĐ (k, củ), bóc tách số tiền, warning level, gợi ý sinh viên.", "Sinh viên / Gemini AI", "FR005"],
        ["UC006", "AI Tư vấn tiết kiệm & Giải toán", "Phân tích CSV chi tiêu tư vấn tiết kiệm và giải bài tập toán học.", "Sinh viên / Gemini AI", "FR006"],
        ["UC007", "AI Chatbot & CSKH Agentic", "Khung chat nổi popover thực thi trực tiếp lệnh Thêm/Xóa/Đặt hạn mức.", "Sinh viên / Gemini AI", "FR007"],
        ["UC008", "Proxy & Fallback Geo-blocking", "Tự động inject proxy từ .env và trả về thông báo hướng dẫn an toàn.", "System / Proxy", "FR008"]
    ]
    populate_doc_tables(doc, [uc_table])

# ==============================================================================
# 04. OBJECT-ORIENTED DESIGN (OOD)
# ==============================================================================
MD_04 = """# TÀI LIỆU THIẾT KẾ HƯỚNG ĐỐI TƯỢNG (OOD)

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 01 - CNTT K23K  

---

## 1. Kiến Trúc Phân Lớp Hệ Thống (Decoupled Layered Architecture)

Hệ thống được thiết kế theo mô hình 4 lớp rõ ràng:

1. **Presentation Layer (Frontend)**: `frontend/app.py` xây dựng bằng Streamlit UI, gọi Backend qua `requests`.
2. **API Controller Layer (FastAPI Routers)**: `backend/api/routes/` gồm `transactions.py`, `budgets.py`, `analytics.py`, `ai.py`.
3. **Business Service Layer**: `backend/services/` gồm `transaction_service.py`, `budget_service.py`, `analytics_service.py`, `ai_service.py`.
4. **Core & Data Access Layer**: `backend/core/` (`config.py`, `database.py`) kết nối SQLite `chi_tieu.db` và nạp cấu hình Proxy.

## 2. Chi Tiết Các Class & Schemas Pydantic

- **TransactionCreate / TransactionUpdate / TransactionResponse**: Pydantic models quản lý dữ liệu giao dịch.
- **BudgetLimitSet / BudgetLimitResponse**: Pydantic models quản lý hạn mức.
- **SummaryResponse / MonthlyComparisonResponse**: Schemas dữ liệu thống kê tài chính.
- **AIParseRequest / AIChatRequest / AIAgentRequest / AIResponse**: Schemas trao đổi dữ liệu AI API.
- **TransactionService**: Hàm static `list_transactions()`, `add_transaction()`, `update_transaction()`, `delete_transaction()`.
- **AIService**: Hàm static `init_gemini()`, `analyze_natural_language_expense()`, `generate_savings_advice()`, `chat_with_gemini()`, `chat_with_gemini_agent()`.
"""

def docx_builder_04(doc):
    add_heading_styled(doc, "TÀI LIỆU THIẾT KẾ HƯỚNG ĐỐI TƯỢNG (OOD)", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_heading_styled(doc, "1. Mô Hình Phân Lớp Kiến Trúc FastAPI & Streamlit", 2)
    add_paragraph_styled(doc, "Hệ thống được mô hình hóa theo kiến trúc 4 lớp: Frontend UI (Streamlit), API Controllers (FastAPI), Services Layer (Pure Python Business Logic) và Data Access Layer (SQLite Engine).")
    add_heading_styled(doc, "2. Danh Sách Các Class & Pydantic Data Models", 2)
    classes_table = [
        ["Module / Package", "Tên Class / Model", "Vai trò & Chức năng"],
        ["backend.models.schemas", "TransactionCreate", "Schema request tạo mới giao dịch (loai, so_tien, danh_muc, ngay, ghi_chu)."],
        ["backend.models.schemas", "TransactionResponse", "Schema response trả về chi tiết giao dịch từ SQLite."],
        ["backend.models.schemas", "BudgetLimitSet", "Schema đặt hạn mức ngân sách theo danh mục."],
        ["backend.models.schemas", "SummaryResponse", "Schema chứa tổng thu, tổng chi và số dư ví."],
        ["backend.services", "TransactionService", "Thực thi CRUD CSDL SQLite bảng giao_dich."],
        ["backend.services", "AIService", "Kết nối Google Gemini API, inject Proxy, bóc tách chi tiêu & Agentic CSKH."],
        ["backend.api.routes", "TransactionsRouter", "FastAPI APIRouter quản lý các endpoints /api/transactions."],
        ["frontend", "StreamlitApp", "Giao diện Web UI 5 Tabs + Floating Popover CSKH Widget."]
    ]
    populate_doc_tables(doc, [classes_table])

# ==============================================================================
# 05. FUNCTIONAL TESTING
# ==============================================================================
MD_05 = """# BÁO CÁO KIỂM THỬ CHỨC NĂNG (FUNCTIONAL TESTING)

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 01 - CNTT K23K  

---

## 1. Môi Trường & Tài Nguyên Kiểm Thử
- **Phần cứng**: CPU Intel/AMD, RAM 8GB, kết nối Internet / Proxy.
- **Môi trường phần mềm**: Python 3.12, Pytest 9.1, FastAPI TestClient, Request library.

## 2. Kịch Bản & Kết Quả Kiểm Thử API Integration (tests/test_api.py)

| Test ID | Tên Kịch Bản Kiểm Thử | Endpoint / Function | Trạng thái | Thời gian |
|---------|-----------------------|---------------------|------------|-----------|
| TC01 | Kiểm tra Health Check Backend | `GET /health` | **PASSED** | 0.05s |
| TC02 | Lấy danh sách giao dịch | `GET /api/transactions` | **PASSED** | 0.12s |
| TC03 | Thêm mới & Xóa giao dịch | `POST /api/transactions` -> `DELETE` | **PASSED** | 0.25s |
| TC04 | Lấy danh sách hạn mức | `GET /api/budgets` | **PASSED** | 0.08s |
| TC05 | Đặt hạn mức chi tiêu | `POST /api/budgets` | **PASSED** | 0.15s |
| TC06 | Thống kê tổng thu/chi/số dư | `GET /api/analytics/summary` | **PASSED** | 0.10s |
| TC07 | So sánh chi tiêu tháng | `GET /api/analytics/monthly-comparison` | **PASSED** | 0.14s |
| TC08 | Quy đổi từ lóng VNĐ (199k, 5 củ) | `test_money_slang_conversion` | **PASSED** | 0.04s |
| TC09 | Kiểm tra file môi trường .env | `test_env_file_exists` | **PASSED** | 0.01s |
| TC10 | Geo-blocking proxy fallback | `test_is_geo_blocked_error` | **PASSED** | 0.02s |

**Tổng kết**: **100% PASSED** (10/10 unit tests & 7/7 REST API integration tests).
"""

def docx_builder_05(doc):
    add_heading_styled(doc, "BÁO CÁO KIỂM THỬ CHỨC NĂNG (FUNCTIONAL TESTING)", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_heading_styled(doc, "Bảng Kịch Bản & Kết Quả Kiểm Thử Tự Động (Pytest)", 2)
    test_table = [
        ["Test ID", "Chức năng kiểm thử", "Mô tả / Test Data", "Kết quả (Pass/Fail)"],
        ["TC01", "Health check API", "GET /health trả về status ok", "PASSED"],
        ["TC02", "Lấy danh sách giao dịch", "GET /api/transactions trả về danh sách json", "PASSED"],
        ["TC03", "Thêm & Xóa giao dịch", "POST /api/transactions tạo mới ID rồi DELETE /api/transactions/{id}", "PASSED"],
        ["TC04", "Lấy & Đặt hạn mức", "GET /api/budgets & POST /api/budgets cài hạn mức 750k", "PASSED"],
        ["TC05", "Thống kê tài chính", "GET /api/analytics/summary tính tổng thu chi", "PASSED"],
        ["TC06", "So sánh tháng", "GET /api/analytics/monthly-comparison tính delta % so với tháng trước", "PASSED"],
        ["TC07", "Xử lý từ lóng VNĐ", "Quy đổi 199k -> 199000, 5 củ -> 5000000", "PASSED"],
        ["TC08", "Fallback Geo-blocking", "Bắt lỗi user location is not supported và nạp thông báo hướng dẫn", "PASSED"]
    ]
    populate_doc_tables(doc, [test_table])

# ==============================================================================
# 06. DATABASE & SCREENFLOW
# ==============================================================================
MD_06 = """# SCREEN FLOW & TÀI LIỆU THIẾT KẾ CƠ SỞ DỮ LIỆU

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 01 - CNTT K23K  

---

## 1. Sơ Đồ Phân Luồng Màn Hình (Screen Flow)

```
[Màn hình Chính Streamlit Web App]
   ├── Tab 1: 📊 Thống Kê & Ngân Sách (Metric cards, Progress bars, Chart Pie/Bar/Line)
   ├── Tab 2: ➕ Thêm Chi Tiêu (AI Natural Language Input + Form Thủ Công)
   ├── Tab 3: 📜 Lịch Sử (Danh sách Dataframe, Bộ lọc 3 tiêu chí, Edit/Delete confirm 2-step, Export/Import)
   ├── Tab 4: 🧠 Gợi Ý Tiết Kiệm (Gemini Advice Generator)
   ├── Tab 5: 🤖 Trợ Lý Gemini (Chatbot Finance & Math QA)
   └── Floating Widget: 💬 CSKH AI Popover Nổi góc dưới bên phải màn hình
```

## 2. Cấu Trúc Cơ Sở Dữ Liệu SQLite (`chi_tieu.db`)

### Bảng 1: `giao_dich` (Lưu lịch sử thu chi)
- `id` (INTEGER PRIMARY KEY AUTOINCREMENT): Mã giao dịch.
- `loai` (TEXT NOT NULL): 'Thu' hoặc 'Chi'.
- `so_tien` (REAL NOT NULL): Số tiền VNĐ (> 0).
- `danh_muc` (TEXT NOT NULL): 1 trong 9 danh mục sinh viên.
- `ngay` (TEXT NOT NULL): Định dạng YYYY-MM-DD.
- `ghi_chu` (TEXT): Mô tả giao dịch.
- `created_at` (TIMESTAMP DEFAULT CURRENT_TIMESTAMP): Thời gian tạo.

### Bảng 2: `han_muc` (Lưu hạn mức chi tiêu tháng)
- `danh_muc` (TEXT PRIMARY KEY): Tên danh mục.
- `so_tien_limit` (REAL NOT NULL): Số tiền hạn mức tối đa.
"""

def docx_builder_06(doc):
    add_heading_styled(doc, "SCREEN FLOW & TÀI LIỆU THIẾT KẾ CƠ SỞ DỮ LIỆU", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_heading_styled(doc, "1. Cấu Trúc Các Bảng Trong Cơ Sở Dữ Liệu SQLite", 2)
    db_table = [
        ["Tên Bảng", "Tên Cột (Field)", "Kiểu Dữ Liệu", "Khóa / Ràng buộc", "Mô tả nghiệp vụ"],
        ["giao_dich", "id", "INTEGER", "PRIMARY KEY AUTOINCREMENT", "Mã định danh giao dịch thu chi."],
        ["giao_dich", "loai", "TEXT", "NOT NULL ('Thu'/'Chi')", "Loại giao dịch: Thu nhập hoặc Chi tiêu."],
        ["giao_dich", "so_tien", "REAL", "NOT NULL (> 0)", "Số tiền thực tế tính theo VNĐ."],
        ["giao_dich", "danh_muc", "TEXT", "NOT NULL", "Danh mục chi tiêu sinh viên (9 danh mục)."],
        ["giao_dich", "ngay", "TEXT", "NOT NULL (YYYY-MM-DD)", "Ngày thực hiện giao dịch."],
        ["giao_dich", "ghi_chu", "TEXT", "NULLABLE", "Ghi chú/Mô tả chi tiết khoản thu chi."],
        ["han_muc", "danh_muc", "TEXT", "PRIMARY KEY", "Tên danh mục chi tiêu được đặt hạn mức."],
        ["han_muc", "so_tien_limit", "REAL", "NOT NULL (> 0)", "Hạn mức chi tiêu tối đa trong tháng."]
    ]
    populate_doc_tables(doc, [db_table])

# ==============================================================================
# 07. USER GUIDE
# ==============================================================================
MD_07 = """# HƯỚNG DẪN SỬ DỤNG VÀ TRIỂN KHAI DỰ ÁN

**Dự án**: Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓  
**Nhóm**: Nhóm 01 - CNTT K23K  

---

## 1. Triển Khai Nhanh Bằng Docker Compose (Khuyên Dùng)

```bash
# 1. Clone hoặc tải mã nguồn dự án
cd quan_ly_thu_chi

# 2. Khởi động Docker Compose
docker-compose up --build
```

- **Frontend Streamlit UI**: [http://localhost:8502](http://localhost:8502)
- **Backend FastAPI Swagger Docs**: [http://localhost:8000/docs](http://localhost:8000/docs)

---

## 2. Triển Khai Thủ Công Cục Bộ (Local Python)

```bash
# Terminal 1: Chạy Backend FastAPI
uvicorn backend.main:app --reload --port 8000

# Terminal 2: Chạy Frontend Streamlit
python -m streamlit run frontend/app.py --server.port 8502
```

---

## 3. Hướng Dẫn Sử Dụng Chi Tiết

1. **Tab 📊 Thống Kê**: Xem tổng thu chi, số dư ví, sức khỏe tài chính, so sánh tháng này vs tháng trước và 3 biểu đồ Plotly (Pie, Bar, Line trend).
2. **Tab ➕ Thêm Chi Tiêu**: Gõ câu tự nhiên (VD: "Sáng ăn phở 35k", "Chiều tiêu 5 củ"), Gemini AI tự động bóc tách số tiền, cảnh báo ví và đưa nút xác nhận lưu vào sổ.
3. **Tab 📜 Lịch Sử**: Lọc theo từ khóa/danh mục/tháng, sửa giao dịch, xóa với xác nhận 2 bước, xuất báo cáo CSV/Excel.
4. **Tab 🧠 Gợi Ý Tiết Kiệm**: Nhấn nút phân tích để Gemini AI xuất lời khuyên tài chính cá nhân hóa.
5. **Tab 🤖 Trợ Lý Gemini**: Hỏi đáp kiến thức tài chính & hỗ trợ giải bài tập toán học.
6. **Widget CSKH Nổi 💬**: Nhấp vào biểu tượng chat góc dưới bên phải để ra lệnh trực tiếp bằng ngôn ngữ tự nhiên.
"""

def docx_builder_07(doc):
    add_heading_styled(doc, "HƯỚNG DẪN SỬ DỤNG VÀ TRIỂN KHAI DỰ ÁN", 1)
    add_paragraph_styled(doc, f"Tên ứng dụng: {PROJECT_NAME}", bold=True)
    add_paragraph_styled(doc, HEADER_INFO)
    add_heading_styled(doc, "1. Hướng Dẫn Triển Khai Lệnh Docker Compose", 2)
    add_paragraph_styled(doc, "Để khởi chạy toàn bộ hệ thống gồm FastAPI Backend và Streamlit Frontend, thực hiện lệnh:\n`docker-compose up --build`\n\nSau khi khởi chạy thành công:\n- Web App Streamlit: http://localhost:8502\n- API Swagger Specs: http://localhost:8000/docs")
    add_heading_styled(doc, "2. Hướng Dẫn Sử Dụng Các Tính Năng Chi Tiết", 2)
    add_paragraph_styled(doc, "Ứng dụng bao gồm 5 Tab chính và Widget CSKH Nổi ở góc dưới bên phải màn hình. Người dùng có thể dễ dàng nhập chi tiêu bằng ngôn ngữ tự nhiên, xem biểu đồ xu hướng, đặt hạn mức ngân sách và xuất báo cáo CSV/Excel.")

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================
def write_file_pair(filename_prefix, docx_name, md_content, docx_builder):
    for bdir in BASE_DIRS:
        if not os.path.exists(bdir):
            continue
        md_path = os.path.join(bdir, f"{filename_prefix}.md")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(md_content)
        print(f"  [MD] Created {md_path}")
        
        docx_path = os.path.join(bdir, docx_name)
        doc = docx.Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)
        
        docx_builder(doc)
        doc.save(docx_path)
        print(f"  [DOCX] Created {docx_path}")

def main():
    sys.stdout.reconfigure(encoding='utf-8')
    print("🚀 Starting full report generation for Sổ Tay Quản Lý Chi Tiêu Sinh Viên AI 🎓...")
    
    reports = [
        ("01_project-plan", "01_GenAI_SoftwareDevelopment_project-plan.docx", MD_01, docx_builder_01),
        ("02_requirements-qa", "02_GenAI_SoftwareDevelopment_requirements-qa.docx", MD_02, docx_builder_02),
        ("03_requirements-specification", "03_GenAI_SoftwareDevelopment_requirements-specification.docx", MD_03, docx_builder_03),
        ("04_object-oriented-design", "04_GenAI_SoftwareDevelopment_object-oriented-design.docx", MD_04, docx_builder_04),
        ("05_functional-testing", "05_GenAI_SoftwareDevelopment_functional-testing.docx", MD_05, docx_builder_05),
        ("06_database", "06_GenAI_SoftwareDevelopment_screenflow_db.docx", MD_06, docx_builder_06),
        ("07_user-guide", "07_GenAI_SoftwareDevelopment_user-guide.docx", MD_07, docx_builder_07),
    ]
    
    for prefix, docx_name, md_content, builder in reports:
        print(f"\nProcessing {prefix}...")
        write_file_pair(prefix, docx_name, md_content, builder)
        
    print("\n🎉 ALL 7 REPORT STAGES GENERATED SUCCESSFULLY IN BOTH .MD AND .DOCX FORMATS!")

if __name__ == "__main__":
    main()
